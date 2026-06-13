from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def render_report_docx(
    template_path: str,
    output_path: Path,
    metrics: Dict[str, Any],
    profile: Dict[str, Any],
    report_draft: Dict[str, Any],
    chart_paths: List[str],
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(template_path)
    document._body.clear_content()
    configure_document_styles(document)

    document.add_heading("停车明细分析报告", level=1)
    document.add_paragraph(report_draft["executive_summary"])

    document.add_heading("一、核心指标", level=2)
    metric_rows = [
        ("总交易笔数", f"{metrics['total_transactions']} 笔"),
        ("应收总金额", f"{metrics['receivable_amount']:.2f} 元"),
        ("实收总金额", f"{metrics['collected_amount']:.2f} 元"),
        ("实际抵扣总额", f"{metrics['actual_deduct_amount']:.2f} 元"),
        ("实收率", f"{metrics['collection_rate']:.1f}%"),
        ("主要支付方式", f"{metrics['top_payment_method']}（{metrics['top_payment_method_count']} 笔）"),
    ]
    add_kv_table(document, ["指标", "结果"], metric_rows, (2.2, 3.9))

    document.add_heading("二、交易画像", level=2)
    document.add_paragraph(
        f"平均停车时长 {profile['duration']['average_hours']} 小时，"
        f"最长停车时长 {profile['duration']['max_hours']} 小时。"
    )
    document.add_heading("支付渠道", level=3)
    add_kv_table(
        document,
        ["渠道", "交易笔数"],
        [(item["name"], f"{item['count']} 笔") for item in profile["payment_channel_counts"]],
        (3.0, 3.1),
    )
    document.add_heading("支付方式", level=3)
    add_kv_table(
        document,
        ["支付方式", "交易笔数"],
        [(item["name"], f"{item['count']} 笔") for item in profile["payment_method_counts"]],
        (3.0, 3.1),
    )
    document.add_heading("停车时长分布", level=3)
    add_kv_table(
        document,
        ["时长区间（小时）", "交易笔数"],
        [(item["range"], f"{item['count']} 笔") for item in profile["duration"]["bins"]],
        (3.0, 3.1),
    )
    for chart_path in chart_paths:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(chart_path, width=Inches(5.5))

    document.add_heading("三、补充观察", level=2)
    for observation in report_draft["observations"]:
        add_paragraph_with_optional_style(document, observation, "List Bullet", "· ")

    document.add_heading("四、管理建议", level=2)
    for index, recommendation in enumerate(report_draft["recommendations"], start=1):
        add_paragraph_with_optional_style(document, recommendation, "List Number", f"{index}. ")

    document.save(output_path)
    return output_path


def configure_document_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_kv_table(
    document: Document,
    headers: Sequence[str],
    rows: Iterable[Tuple[str, str]],
    widths: Tuple[float, float],
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_borders(table)
    set_table_cell_margins(table)

    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, widths[index])
        set_cell_text(cell, text, bold=True)
        set_cell_shading(cell, "DCEADF")

    for name, value in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])
        set_cell_text(cells[0], name)
        set_cell_text(cells[1], value)

    document.add_paragraph()


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "7A8A92")


def set_table_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for edge in ["top", "left", "bottom", "right"]:
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), "120")
        element.set(qn("w:type"), "dxa")


def add_paragraph_with_optional_style(document: Document, text: str, style: str, fallback_prefix: str) -> None:
    if has_paragraph_style(document, style):
        document.add_paragraph(text, style=style)
    else:
        document.add_paragraph(f"{fallback_prefix}{text}")


def has_paragraph_style(document: Document, style: str) -> bool:
    return any(item.name == style and item.type == WD_STYLE_TYPE.PARAGRAPH for item in document.styles)
